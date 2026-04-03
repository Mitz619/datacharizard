"""
tasks/resume_builder.py  —  Task 4: AI Resume Customiser
Usage:
  python main.py resume
  python main.py resume --reparse
"""
import os, json, re
from datetime import date
from pathlib import Path

BASE_DIR    = Path(__file__).parent.parent
RESUME_DIR  = BASE_DIR / "resume"
OUTPUT_DIR  = RESUME_DIR / "output"
BASE_RESUME = RESUME_DIR / "base_resume.pdf"

RESUME_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)


# ── Prompts ────────────────────────────────────────────────────────────────────

PARSE_SYSTEM = """You are a resume parser. Extract the resume into structured JSON.
Return ONLY valid JSON, no markdown, exactly this structure:
{
  "name": "FULL NAME",
  "subtitle": "Job Title | Skill | Skill | ...",
  "email": "email",
  "phone": "phone",
  "location": "City, State",
  "linkedin": "linkedin URL or empty string",
  "summary": "full summary paragraph",
  "certifications": [
    {"name": "Cert full name", "credential_id": "ID here"}
  ],
  "skills": ["Skill1", "Skill2"],
  "experience": [
    {
      "dates": "MM/YYYY - MM/YYYY or Present",
      "location": "City, Country",
      "title": "Job Title",
      "company": "Company Name",
      "description": "the italic overview paragraph",
      "bullets": ["bullet point one", "bullet point two"]
    }
  ],
  "education": [
    {
      "dates": "YYYY - YYYY",
      "location": "City, Country",
      "degree": "Degree Name",
      "institution": "University Name"
    }
  ]
}
Preserve ALL bullet points from the original resume exactly."""


CUSTOMISE_SYSTEM = """You are an expert resume writer for data engineering and analytics roles in Australia.

Given a candidate's resume JSON and a target job, produce a highly tailored, crisp, effective version.

CRITICAL RULES:
1. SUBTITLE — set as: "Exact Job Title | TopSkill | TopSkill | TopSkill | TopSkill | TopSkill"
   Use the exact job title from the listing, then the 5 most relevant skills from the job ad

2. SUMMARY — max 3 punchy sentences (under 70 words total):
   - Sentence 1: years of experience + most relevant domain for THIS company
   - Sentence 2: top 2-3 technical skills that match the job listing
   - Sentence 3: one quantified achievement most relevant to this role

3. EXPERIENCE BULLETS — for each role:
   - Most recent role: keep max 8 bullets, most relevant to this job first
   - Other roles: keep max 5 bullets, most relevant first
   - Rephrase bullets to use keywords from the job listing
   - Remove or cut bullets not relevant to this role

4. SKILLS — reorder: matching skills first, then others

5. Keep certifications and education exactly as-is (never change credential IDs)
6. Keep experience in reverse-chronological order (most recent first)
7. Never invent experience or qualifications
8. Return ONLY valid JSON with identical structure — no markdown, no extra keys"""


# ── PDF extraction ─────────────────────────────────────────────────────────────

def extract_pdf_text(pdf_path: Path) -> str:
    try:
        import pdfplumber
    except ImportError:
        os.system("pip install pdfplumber -q")
        import pdfplumber
    text = ""
    with pdfplumber.open(str(pdf_path)) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text += t + "\n"
    return text.strip()


# ── Parse + cache ──────────────────────────────────────────────────────────────

def parse_resume(pdf_path: Path) -> dict:
    from utils.db import get_conn
    from utils.claude_client import ask

    with get_conn() as conn:
        conn.execute("""CREATE TABLE IF NOT EXISTS resume_base (
            id INTEGER PRIMARY KEY, parsed_json TEXT, updated_on TEXT)""")
        row = conn.execute(
            "SELECT parsed_json FROM resume_base LIMIT 1").fetchone()

    if row:
        print("  ✅ Using cached resume (use --reparse to refresh)")
        return json.loads(row[0])

    print("  📄 Extracting PDF text...")
    raw_text = extract_pdf_text(pdf_path)
    print("  🤖 Parsing with Claude...")
    raw = ask(PARSE_SYSTEM, f"Parse this resume:\n\n{raw_text}", max_tokens=4000)
    raw = re.sub(r"```json|```", "", raw).strip()
    parsed = json.loads(raw)

    with get_conn() as conn:
        conn.execute("DELETE FROM resume_base")
        conn.execute(
            "INSERT INTO resume_base (parsed_json, updated_on) VALUES (?,?)",
            (json.dumps(parsed), str(date.today())))

    print(f"  ✅ Parsed: {parsed.get('name')} — "
          f"{len(parsed.get('experience',[]))} roles, "
          f"{len(parsed.get('skills',[]))} skills")
    return parsed


def reparse_resume(pdf_path: Path) -> dict:
    from utils.db import get_conn
    with get_conn() as conn:
        conn.execute("DELETE FROM resume_base")
    return parse_resume(pdf_path)


# ── Job picker ─────────────────────────────────────────────────────────────────

def pick_job() -> dict | None:
    from utils.db import get_conn
    with get_conn() as conn:
        jobs = conn.execute("""
            SELECT id, title, company, location, source, url, skills
            FROM jobs ORDER BY scraped_on DESC LIMIT 30
        """).fetchall()

    if not jobs:
        print("  ❌ No jobs in DB. Run: python main.py jobs")
        return None

    print(f"\n  {'#':<4} {'Title':<42} {'Company':<25} {'Location':<20} Source")
    print("  " + "─" * 100)
    for i, j in enumerate(jobs, 1):
        print(f"  {i:<4} {j['title'][:40]:<42} {j['company'][:23]:<25} "
              f"{j['location'][:18]:<20} {j['source']}")
    print()

    try:
        choice = int(input("  Enter job number: ").strip())
        if 1 <= choice <= len(jobs):
            job = dict(jobs[choice - 1])
            job["skills"] = json.loads(job.get("skills") or "[]")
            return job
    except (ValueError, IndexError):
        pass
    print("  ❌ Invalid selection")
    return None


# ── Customise ──────────────────────────────────────────────────────────────────

def customise_resume(base: dict, job: dict) -> dict:
    from utils.claude_client import ask
    prompt = f"""Target job:
Title: {job['title']}
Company: {job['company']}
Location: {job['location']}
Skills mentioned: {job.get('skills', [])}

Candidate resume JSON:
{json.dumps(base, indent=2)}

Produce a tailored resume JSON for {job['title']} at {job['company']}."""

    print(f"  🤖 Tailoring for {job['title']} at {job['company']}...")
    raw = ask(CUSTOMISE_SYSTEM, prompt, max_tokens=4000)
    raw = re.sub(r"```json|```", "", raw).strip()
    result = json.loads(raw)
    print("  ✅ Resume tailored!")
    return result


# ── Sort helper ────────────────────────────────────────────────────────────────

def sort_reverse_chron(entries: list) -> list:
    def year_key(e):
        dates = e.get("dates", "")
        if "present" in dates.lower():
            return 9999
        nums = re.findall(r"\d{4}", dates)
        return int(nums[0]) if nums else 0
    return sorted(entries, key=year_key, reverse=True)


# ── PDF generation ─────────────────────────────────────────────────────────────

def generate_pdf(resume: dict, output_path: Path):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                     HRFlowable, Table, TableStyle, KeepTogether)
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY

    # ── Colours ────────────────────────────────────────────────────
    BLUE        = colors.HexColor("#2E74B5")
    BLACK       = colors.HexColor("#000000")
    DARK        = colors.HexColor("#1a1a1a")
    GRAY        = colors.HexColor("#666666")
    BORDER_GRAY = colors.HexColor("#CCCCCC")

    # ── Page metrics ───────────────────────────────────────────────
    PAGE_W, _ = A4
    LM, RM    = 18*mm, 18*mm
    CW        = PAGE_W - LM - RM
    DATE_W    = 30*mm
    BULLET_W  =  7*mm
    CONT_W    = CW - DATE_W - BULLET_W

    def sty(name, **kw):
        return ParagraphStyle(name, **kw)

    # ── Styles ─────────────────────────────────────────────────────
    S_NAME   = sty("Name", fontSize=22, fontName="Helvetica-Bold",
                   textColor=BLACK, spaceAfter=2, leading=26)
    S_SUB    = sty("Sub",  fontSize=10, fontName="Helvetica-Bold",
                   textColor=BLUE, spaceAfter=4, leading=14)
    S_CON    = sty("Con",  fontSize=9,  fontName="Helvetica",
                   textColor=GRAY, spaceAfter=1, leading=13)
    # keepWithNext prevents section heading orphaned at page bottom
    S_SEC    = sty("Sec",  fontSize=11, fontName="Helvetica-Bold",
                   textColor=BLACK, spaceBefore=6, spaceAfter=2,
                   leading=14, keepWithNext=1)
    S_BODY_J = sty("BodJ", fontSize=9,  fontName="Helvetica",
                   textColor=DARK, leading=14, spaceAfter=3,
                   alignment=TA_JUSTIFY)
    S_DATE   = sty("Dat",  fontSize=9,  fontName="Helvetica-Bold",
                   textColor=DARK, leading=14)
    S_LOC    = sty("Loc",  fontSize=9,  fontName="Helvetica",
                   textColor=GRAY, leading=13)
    S_ETIT   = sty("ETit", fontSize=10, fontName="Helvetica",
                   textColor=DARK, leading=14, spaceAfter=1)
    S_ECO    = sty("ECo",  fontSize=10, fontName="Helvetica-Bold",
                   textColor=BLUE, leading=13, spaceAfter=2)
    S_EITA   = sty("EIta", fontSize=9,  fontName="Helvetica-Oblique",
                   textColor=DARK, leading=13, spaceAfter=3,
                   alignment=TA_JUSTIFY)
    S_BUL    = sty("Bul",  fontSize=9,  fontName="Helvetica",
                   textColor=DARK, leading=14, spaceAfter=0,
                   leftIndent=4)
    S_DOT    = sty("Dot",  fontSize=11, fontName="Helvetica",
                   textColor=DARK, leading=14, alignment=TA_CENTER)
    S_CERTN  = sty("CerN", fontSize=9,  fontName="Helvetica-Bold",
                   textColor=DARK, leading=13, spaceAfter=1)
    S_CERID  = sty("CerI", fontSize=8,  fontName="Helvetica",
                   textColor=GRAY, leading=12, spaceAfter=5)
    S_SKILL  = sty("Ski",  fontSize=9,  fontName="Helvetica",
                   textColor=DARK, leading=14)
    S_REF    = sty("Ref",  fontSize=9,  fontName="Helvetica-Bold",
                   textColor=DARK, leading=13)
    S_EMPTY  = sty("Emp",  fontSize=9,  fontName="Helvetica",
                   textColor=DARK, leading=14)

    # ── Document ───────────────────────────────────────────────────
    doc = SimpleDocTemplate(
        str(output_path), pagesize=A4,
        leftMargin=LM, rightMargin=RM,
        topMargin=14*mm, bottomMargin=14*mm,
    )

    def hr():
        return HRFlowable(width="100%", thickness=0.5,
                          color=BORDER_GRAY, spaceAfter=5, spaceBefore=1)

    def section(title):
        # keepWithNext=1 on S_SEC means heading sticks to next element
        return [Paragraph(title.upper(), S_SEC), hr()]

    def make_exp_table(exp: dict) -> Table:
        """
        Multi-row table for one experience entry.
        Row 0: date/loc | • | title + company + description
        Row N: ""       | "" | bullet N
        Last:  ""       | "" | spacer

        Multi-row tables split naturally across pages between bullets,
        eliminating the whitespace problem caused by KeepTogether.
        The vertical line runs the full height via LINEBEFORE on col 2.
        """
        rows = []

        # Header row
        left_header = [
            Paragraph(exp.get("dates", ""),    S_DATE),
            Paragraph(exp.get("location", ""), S_LOC),
        ]
        right_header = [
            Paragraph(exp.get("title", ""),   S_ETIT),
            Paragraph(exp.get("company", ""), S_ECO),
        ]
        if exp.get("description"):
            right_header.append(Paragraph(exp["description"], S_EITA))

        rows.append([left_header, Paragraph("•", S_DOT), right_header])

        # One row per bullet
        for b in exp.get("bullets", []):
            b_html = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", b)
            rows.append([
                Paragraph("", S_EMPTY),
                Paragraph("", S_EMPTY),
                Paragraph(f"· {b_html}", S_BUL),
            ])

        # Spacer row at end of entry
        rows.append([Spacer(1, 8), Spacer(1, 8), Spacer(1, 8)])

        t = Table(rows, colWidths=[DATE_W, BULLET_W, CONT_W])
        t.setStyle(TableStyle([
            ("VALIGN",        (0, 0), (-1, -1), "TOP"),
            ("TOPPADDING",    (0, 0), (-1, -1), 1),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
            ("LEFTPADDING",   (0, 0), (-1, -1), 0),
            ("RIGHTPADDING",  (0, 0), (-1, -1), 0),
            # Vertical line full height of entry
            ("LINEBEFORE",    (2, 0), (2, -1), 0.5, BORDER_GRAY),
            ("LEFTPADDING",   (2, 0), (2, -1), 8),
            ("RIGHTPADDING",  (0, 0), (0, -1), 4),
            # Spacer row: remove padding
            ("TOPPADDING",    (0, -1), (-1, -1), 0),
            ("BOTTOMPADDING", (0, -1), (-1, -1), 0),
        ]))
        return t

    def make_edu_table(edu: dict) -> Table:
        rows = [
            [
                [Paragraph(edu.get("dates", ""),    S_DATE),
                 Paragraph(edu.get("location", ""), S_LOC)],
                Paragraph("•", S_DOT),
                [Paragraph(edu.get("degree", ""),      S_ETIT),
                 Paragraph(edu.get("institution", ""), S_ECO)],
            ],
            [Spacer(1, 6), Spacer(1, 6), Spacer(1, 6)],
        ]
        t = Table(rows, colWidths=[DATE_W, BULLET_W, CONT_W])
        t.setStyle(TableStyle([
            ("VALIGN",        (0,0), (-1,-1), "TOP"),
            ("TOPPADDING",    (0,0), (-1,-1), 1),
            ("BOTTOMPADDING", (0,0), (-1,-1), 1),
            ("LEFTPADDING",   (0,0), (-1,-1), 0),
            ("RIGHTPADDING",  (0,0), (-1,-1), 0),
            ("LINEBEFORE",    (2,0), (2,-1),  0.5, BORDER_GRAY),
            ("LEFTPADDING",   (2,0), (2,-1),  8),
            ("RIGHTPADDING",  (0,0), (0,-1),  4),
            ("TOPPADDING",    (0,-1), (-1,-1), 0),
            ("BOTTOMPADDING", (0,-1), (-1,-1), 0),
        ]))
        return t

    story = []

    # ── HEADER ────────────────────────────────────────────────────
    story.append(Paragraph(resume.get("name", ""), S_NAME))
    if resume.get("subtitle"):
        story.append(Paragraph(resume["subtitle"], S_SUB))

    contact_parts = []
    if resume.get("phone"):    contact_parts.append(resume["phone"])
    if resume.get("email"):    contact_parts.append(resume["email"])
    if resume.get("linkedin"): contact_parts.append(resume["linkedin"])
    if contact_parts:
        story.append(Paragraph("   ·   ".join(contact_parts), S_CON))
    if resume.get("location"):
        story.append(Paragraph(resume["location"], S_CON))

    story.append(Spacer(1, 4))
    story.append(hr())

    # ── SUMMARY ───────────────────────────────────────────────────
    if resume.get("summary"):
        story += section("Summary")
        story.append(Paragraph(resume["summary"], S_BODY_J))
        story.append(Spacer(1, 3))

    # ── CERTIFICATIONS ────────────────────────────────────────────
    certs = resume.get("certifications", [])
    if certs:
        cert_items = section("Certifications")
        col_w = CW / 2
        for i in range(0, len(certs), 2):
            pair = certs[i:i+2]
            cells = []
            for cert in pair:
                if isinstance(cert, dict):
                    n, cid = cert.get("name",""), cert.get("credential_id","")
                else:
                    parts = str(cert).split("Credential ID:")
                    n   = parts[0].strip()
                    cid = parts[1].strip() if len(parts) > 1 else ""
                cell = [Paragraph(n, S_CERTN)]
                if cid:
                    cell.append(Paragraph(f"Credential ID: {cid}", S_CERID))
                cells.append(cell)
            while len(cells) < 2:
                cells.append([Paragraph("", S_EMPTY)])
            t = Table([cells], colWidths=[col_w, col_w])
            t.setStyle(TableStyle([
                ("VALIGN",        (0,0),(-1,-1),"TOP"),
                ("LEFTPADDING",   (0,0),(-1,-1),0),
                ("RIGHTPADDING",  (0,0),(-1,-1),6),
                ("TOPPADDING",    (0,0),(-1,-1),0),
                ("BOTTOMPADDING", (0,0),(-1,-1),0),
            ]))
            cert_items.append(t)
        story.append(KeepTogether(cert_items))
        story.append(Spacer(1, 3))

    # ── SKILLS — plain text rows with · separator ──────────────────
    skills = resume.get("skills", [])
    if skills:
        skill_items = section("Skills")
        PER_ROW = 8
        rows = [skills[i:i+PER_ROW] for i in range(0, len(skills), PER_ROW)]
        for row in rows:
            skill_items.append(Paragraph("   ·   ".join(row), S_SKILL))
            skill_items.append(Spacer(1, 3))
        story.append(KeepTogether(skill_items))
        story.append(Spacer(1, 3))

    # ── EXPERIENCE — multi-row tables, flow naturally across pages ─
    experience = resume.get("experience", [])
    if experience:
        exp_heading = section("Experience")
        for idx, exp in enumerate(experience):
            t = make_exp_table(exp)
            if idx == 0:
                # Glue heading to first job's header row only
                story.append(KeepTogether(exp_heading + [t]))
            else:
                story.append(t)

    # ── EDUCATION ─────────────────────────────────────────────────
    education = resume.get("education", [])
    if education:
        edu_heading = section("Education")
        for idx, edu in enumerate(education):
            t = make_edu_table(edu)
            if idx == 0:
                story.append(KeepTogether(edu_heading + [t]))
            else:
                story.append(t)

    # ── REFERENCE ─────────────────────────────────────────────────
    ref_block = section("Reference")
    ref_block.append(Paragraph("Available upon request", S_REF))
    story.append(KeepTogether(ref_block))

    doc.build(story)
    print(f"  📄 Saved: {output_path.name}")




# ── Word document generation ───────────────────────────────────────────────────

def generate_docx(resume: dict, output_path: Path):
    """
    Generate an editable Word document matching the resume design.
    Uses python-docx. Install: pip install python-docx
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import copy
    except ImportError:
        import os
        os.system("pip install python-docx -q")
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import copy

    BLUE  = RGBColor(0x2E, 0x74, 0xB5)
    BLACK = RGBColor(0x00, 0x00, 0x00)
    DARK  = RGBColor(0x1a, 0x1a, 0x1a)
    GRAY  = RGBColor(0x66, 0x66, 0x66)

    doc = Document()

    # ── Page margins ────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(1.4)
        section.bottom_margin = Cm(1.4)
        section.left_margin   = Cm(1.8)
        section.right_margin  = Cm(1.8)

    # ── Helper: paragraph formatting ────────────────────────────────
    def para(text="", bold=False, italic=False, size=10,
             colour=None, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=0, space_after=4, keep_with_next=False):
        p = doc.add_paragraph()
        p.alignment = align
        pf = p.paragraph_format
        pf.space_before = Pt(space_before)
        pf.space_after  = Pt(space_after)
        if keep_with_next:
            pf.keep_with_next = True
        if text:
            run = p.add_run(text)
            run.bold   = bold
            run.italic = italic
            run.font.size  = Pt(size)
            run.font.color.rgb = colour or DARK
        return p

    def add_run(p, text, bold=False, italic=False,
                size=9, colour=None):
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size  = Pt(size)
        run.font.color.rgb = colour or DARK
        return run

    def hr_border(p):
        """Add bottom border to paragraph (acts as horizontal rule)."""
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"),   "single")
        bottom.set(qn("w:sz"),    "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "CCCCCC")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def section_heading(title):
        p = para(title.upper(), bold=True, size=11,
                 colour=BLACK, space_before=6, space_after=2,
                 keep_with_next=True)
        hr_border(p)
        return p

    def timeline_table(dates, location, title, company,
                       description="", bullets=None):
        """
        3-column table: date/loc | • | content
        Left border on content column mimics the vertical timeline line.
        """
        bullets = bullets or []

        # Col widths in inches: date=1.18, bullet=0.28, content=rest
        DATE_IN   = 1.18
        BULLET_IN = 0.28
        PAGE_IN   = 5.95   # A4 content width at these margins
        CONT_IN   = PAGE_IN - DATE_IN - BULLET_IN

        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"

        # Remove all borders first
        def clear_borders(cell):
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBdr = OxmlElement("w:tcBdr")
            for side in ("top","left","bottom","right","insideH","insideV"):
                el = OxmlElement(f"w:{side}")
                el.set(qn("w:val"),   "none")
                el.set(qn("w:sz"),    "0")
                el.set(qn("w:space"), "0")
                el.set(qn("w:color"), "auto")
                tcBdr.append(el)
            tcPr.append(tcBdr)

        # Set column widths
        for i, width in enumerate([DATE_IN, BULLET_IN, CONT_IN]):
            tbl.columns[i].width = Inches(width)

        row = tbl.rows[0]
        cells = row.cells

        # Remove borders + set widths on each cell
        for i, (cell, w) in enumerate(zip(cells, [DATE_IN, BULLET_IN, CONT_IN])):
            clear_borders(cell)
            cell.width = Inches(w)

        # Left border on content cell (vertical timeline line)
        tc   = cells[2]._tc
        tcPr = tc.get_or_add_tcPr()
        tcBdr = OxmlElement("w:tcBdr")
        left_el = OxmlElement("w:left")
        left_el.set(qn("w:val"),   "single")
        left_el.set(qn("w:sz"),    "4")
        left_el.set(qn("w:space"), "0")
        left_el.set(qn("w:color"), "CCCCCC")
        tcBdr.append(left_el)
        tcPr.append(tcBdr)

        # Cell padding on content cell
        tcMar = OxmlElement("w:tcMar")
        for side in ("left",):
            m = OxmlElement(f"w:{side}")
            m.set(qn("w:w"),    "113")   # ~8pt
            m.set(qn("w:type"), "dxa")
            tcMar.append(m)
        tcPr.append(tcMar)

        # ── Left cell: dates + location ───────────────────────────
        lc = cells[0]
        lc.paragraphs[0].clear()
        p_date = lc.paragraphs[0]
        p_date.paragraph_format.space_after  = Pt(0)
        p_date.paragraph_format.space_before = Pt(0)
        r = p_date.add_run(dates)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = DARK
        p_loc = lc.add_paragraph(location)
        p_loc.paragraph_format.space_after  = Pt(0)
        p_loc.paragraph_format.space_before = Pt(0)
        p_loc.runs[0].font.size  = Pt(9)
        p_loc.runs[0].font.color.rgb = GRAY

        # ── Middle cell: bullet dot ────────────────────────────────
        mc = cells[1]
        mc.paragraphs[0].clear()
        mc.paragraphs[0].paragraph_format.space_after  = Pt(0)
        mc.paragraphs[0].paragraph_format.space_before = Pt(0)
        mc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = mc.paragraphs[0].add_run("•")
        r.font.size  = Pt(11)
        r.font.color.rgb = DARK

        # ── Right cell: content ────────────────────────────────────
        rc = cells[2]
        rc.paragraphs[0].clear()

        # Title
        p_title = rc.paragraphs[0]
        p_title.paragraph_format.space_after  = Pt(1)
        p_title.paragraph_format.space_before = Pt(0)
        r = p_title.add_run(title)
        r.font.size  = Pt(10)
        r.font.color.rgb = DARK

        # Company
        p_co = rc.add_paragraph()
        p_co.paragraph_format.space_after  = Pt(2)
        p_co.paragraph_format.space_before = Pt(0)
        r = p_co.add_run(company)
        r.bold = True
        r.font.size  = Pt(10)
        r.font.color.rgb = BLUE

        # Description (italic)
        if description:
            p_desc = rc.add_paragraph()
            p_desc.paragraph_format.space_after  = Pt(2)
            p_desc.paragraph_format.space_before = Pt(0)
            p_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r = p_desc.add_run(description)
            r.italic = True
            r.font.size  = Pt(9)
            r.font.color.rgb = DARK

        # Bullets
        for b in bullets:
            # Strip ** markdown bold markers
            b_clean = re.sub(r"\*\*(.+?)\*\*", r"\1", b)
            p_b = rc.add_paragraph()
            p_b.paragraph_format.space_after  = Pt(1)
            p_b.paragraph_format.space_before = Pt(0)
            p_b.paragraph_format.left_indent  = Pt(6)

            # Handle bold parts inside bullets
            parts = re.split(r"(\*\*.+?\*\*)", b)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    r = p_b.add_run("· " + part[2:-2] if p_b.runs == [] else part[2:-2])
                    r.bold = True
                else:
                    text = ("· " + part) if not p_b.runs else part
                    if text.strip():
                        r = p_b.add_run(text)
                r.font.size  = Pt(9)
                r.font.color.rgb = DARK

        # Spacer paragraph at bottom of entry
        p_sp = rc.add_paragraph()
        p_sp.paragraph_format.space_after  = Pt(6)
        p_sp.paragraph_format.space_before = Pt(0)

        return tbl

    # ── Build document ───────────────────────────────────────────────

    # Name
    p = para(resume.get("name",""), bold=True, size=22,
             colour=BLACK, space_before=0, space_after=2)

    # Subtitle
    if resume.get("subtitle"):
        p = para(resume["subtitle"], bold=True, size=10,
                 colour=BLUE, space_before=0, space_after=3)

    # Contact
    contact_parts = []
    if resume.get("phone"):    contact_parts.append(resume["phone"])
    if resume.get("email"):    contact_parts.append(resume["email"])
    if resume.get("linkedin"): contact_parts.append(resume["linkedin"])
    if contact_parts:
        para("   ·   ".join(contact_parts), size=9,
             colour=GRAY, space_after=1)
    if resume.get("location"):
        para(resume["location"], size=9, colour=GRAY, space_after=4)

    # HR line
    p_hr = para("", space_before=0, space_after=4)
    hr_border(p_hr)

    # Summary
    if resume.get("summary"):
        section_heading("Summary")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(resume["summary"])
        r.font.size = Pt(9)
        r.font.color.rgb = DARK

    # Certifications
    certs = resume.get("certifications", [])
    if certs:
        section_heading("Certifications")
        # 2-col table for certs
        for i in range(0, len(certs), 2):
            pair = certs[i:i+2]
            ct = doc.add_table(rows=1, cols=2)
            ct.style = "Table Grid"
            half = Inches(2.975)
            for c in ct.columns:
                c.width = half
            row = ct.rows[0]
            for ci, cert in enumerate(pair):
                if isinstance(cert, dict):
                    n, cid = cert.get("name",""), cert.get("credential_id","")
                else:
                    parts = str(cert).split("Credential ID:")
                    n   = parts[0].strip()
                    cid = parts[1].strip() if len(parts) > 1 else ""
                cell = row.cells[ci]
                cell.width = half
                # Remove borders
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBdr = OxmlElement("w:tcBdr")
                for side in ("top","left","bottom","right","insideH","insideV"):
                    el = OxmlElement(f"w:{side}")
                    el.set(qn("w:val"), "none")
                    tcBdr.append(el)
                tcPr.append(tcBdr)
                cell.paragraphs[0].clear()
                p_n = cell.paragraphs[0]
                p_n.paragraph_format.space_after = Pt(1)
                r = p_n.add_run(n)
                r.bold = True
                r.font.size = Pt(9)
                r.font.color.rgb = DARK
                if cid:
                    p_id = cell.add_paragraph(f"Credential ID: {cid}")
                    p_id.paragraph_format.space_after = Pt(4)
                    p_id.runs[0].font.size = Pt(8)
                    p_id.runs[0].font.color.rgb = GRAY
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Skills
    skills = resume.get("skills", [])
    if skills:
        section_heading("Skills")
        PER_ROW = 8
        rows = [skills[i:i+PER_ROW] for i in range(0, len(skills), PER_ROW)]
        for row in rows:
            p = para("   ·   ".join(row), size=9,
                     colour=DARK, space_after=3)

    # Experience
    experience = resume.get("experience", [])
    if experience:
        section_heading("Experience")
        for exp in experience:
            timeline_table(
                dates=exp.get("dates",""),
                location=exp.get("location",""),
                title=exp.get("title",""),
                company=exp.get("company",""),
                description=exp.get("description",""),
                bullets=exp.get("bullets",[]),
            )
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Education
    education = resume.get("education", [])
    if education:
        section_heading("Education")
        for edu in education:
            timeline_table(
                dates=edu.get("dates",""),
                location=edu.get("location",""),
                title=edu.get("degree",""),
                company=edu.get("institution",""),
            )

    # Reference
    section_heading("Reference")
    p = para("Available upon request", bold=True, size=9, colour=DARK)

    doc.save(str(output_path))
    print(f"  📝 Word doc saved: {output_path.name}")

# ── Entry point ────────────────────────────────────────────────────────────────

def run(reparse: bool = False):
    print("\n" + "═" * 60)
    print("  🔥 DATACHARIZARD — Resume Customiser")
    print("═" * 60)

    if not BASE_RESUME.exists():
        print(f"\n  ❌ Add your resume here: {BASE_RESUME}\n")
        return

    base = reparse_resume(BASE_RESUME) if reparse else parse_resume(BASE_RESUME)

    job = pick_job()
    if not job:
        return

    try:
        customised = customise_resume(base, job)
    except Exception as e:
        print(f"  ⚠️  Tailoring error ({e}) — using base resume")
        customised = base

    # Sort reverse-chronological
    if customised.get("experience"):
        customised["experience"] = sort_reverse_chron(customised["experience"])
    if customised.get("education"):
        customised["education"] = sort_reverse_chron(customised["education"])

    safe_co    = re.sub(r"[^\w\s-]", "", job["company"]).strip().replace(" ", "_")
    safe_title = re.sub(r"[^\w\s-]", "", job["title"]).strip().replace(" ", "_")[:30]
    filename   = f"{safe_co}_{safe_title}_{date.today()}.pdf"
    output     = OUTPUT_DIR / filename

    print("  🖨️  Generating PDF...")
    generate_pdf(customised, output)

    print(f"\n  ✅ PDF done!")
    print(f"     resume/output/{filename}")

    # Ask if they want a Word version
    want_word = input("\n  📝 Generate Word doc for manual editing? (y/n): ").strip().lower()
    if want_word == "y":
        docx_filename = filename.replace(".pdf", ".docx")
        docx_output   = OUTPUT_DIR / docx_filename
        print("  🖨️  Generating Word doc...")
        generate_docx(customised, docx_output)
        print(f"     resume/output/{docx_filename}")
        print("     ✏️  Open in Word or Google Docs to edit freely")

    print(f"\n  Apply → {job.get('url', 'N/A')}\n")


if __name__ == "__main__":
    import sys
    run(reparse="--reparse" in sys.argv)